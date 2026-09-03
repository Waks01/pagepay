"""Study endpoints for Phase 3: AI Exam Prep.

SOW upload → AI parsing → asset generation (MCQ/flashcard/essay) →
ad-or-points gated unlock → streaming study chat.
"""

import logging
import uuid
import asyncio
import io
import re
from datetime import datetime

import httpx

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, Response, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import or_, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.prompts import (
    CHAT_TUTOR_SYSTEM,
    DIAGRAM_GENERATOR,
    ESSAY_ALL_TOPICS_GENERATOR,
    ESSAY_GENERATOR,
    ESSAY_TOPIC_GENERATOR,
    EXAMPLE_GENERATOR,
    FLASHCARD_ALL_TOPICS_GENERATOR,
    FLASHCARD_GENERATOR,
    FLASHCARD_TOPIC_GENERATOR,
    MCQ_ALL_TOPICS_GENERATOR,
    MCQ_GENERATOR,
    MCQ_TOPIC_GENERATOR,
    SOW_PARSER,
    VIDEO_SCRIPT_GENERATOR,
)
from app.ai.router import route_ai
from app.models import (
    AudioUnlock,
    SowUploadJob,
    StudyAsset,
    StudyMaterial,
    StudyProgress,
    StudyTransaction,
    User,
    ReadingSession,
)
from app.database import AsyncSessionLocal, get_db
from app.routers.auth import get_current_user
from app.config import settings
from app.services.sanitize import (
    safe_filename,
    sanitize_for_display,
)
from app.services.tts_hybrid import AUDIO_CACHE_DIR, synthesize_study_audio
from app.schemas import (
    ChatRequest,
    ChatResponse,
    ExampleCheckRequest,
    ExampleCheckResponse,
    ExampleGenerateRequest,
    GenerateAssetRequest,
    GenerateAssetResponse,
    MaterialDetail,
    MaterialSummary,
    MaterialUpdate,
    QuizCompleteRequest,
    QuizCompleteResponse,
    SowUploadJobAccepted,
    SowUploadJobStatus,
    SowUploadRequest,
    SowUploadResponse,
    StudyProgressListResponse,
    StudyProgressResponse,
    StudyProgressUpdate,
    UnlockRequest,
    UnlockResponse,
)

logger = logging.getLogger("uvicorn.error")
router = APIRouter(prefix="/study", tags=["study"])

# Per-route file upload cap. The global RequestSizeLimitMiddleware
# allows up to 1MB, but we tighten that for these routes — a 5MB
# SOW image blows up the base64 payload (2× RAM) and the Gemini
# Vision bill. Documents (PDF/DOCX/TXT) get a slightly larger cap.
MAX_SOW_IMAGE_BYTES: int = 5 * 1024 * 1024   # 5 MB
MAX_SOW_DOC_BYTES: int = 7 * 1024 * 1024     # 7 MB

# Content-type allowlist for the image route. The browser-supplied
# `file.content_type` is attacker-controlled and was previously
# trusted verbatim — restrict it to types Gemini Vision can process.
ALLOWED_IMAGE_TYPES: frozenset[str] = frozenset({
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/webp",
    "image/heic",
    "image/heif",
})

UNLOCK_POINTS_COST = settings.study_unlock_points_cost
VIDEO_UNLOCK_POINTS_COST = settings.study_video_unlock_points_cost


# ── POST /study/sow/upload ──────────────────────────────────────────


@router.post("/sow/upload", response_model=SowUploadResponse, status_code=201)
async def upload_sow(
    payload: SowUploadRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload scheme-of-work text for AI parsing.

    The text is sent to the AI router with the SOW_PARSER prompt.
    On success the parsed JSON is stored alongside the raw input.
    """
    prompt = SOW_PARSER.format(raw_text=payload.text)
    ai_result = await route_ai(prompt, task_type="heavy", db=db)

    parsed = None
    try:
        import json as _json
        parsed = _json.loads(ai_result["response"])
    except Exception:
        logger.error("SOW parser returned non-JSON: %s", ai_result["response"][:200])

    title = (parsed or {}).get("title", "Untitled Material") if isinstance(parsed, dict) else "Untitled Material"

    material = StudyMaterial(
        user_id=current_user.id,
        title=title,
        exam_type=payload.exam_type,
        raw_input=payload.text,
        parsed_structure=_json.dumps(parsed) if parsed else None,
        ai_model_used=ai_result.get("provider"),
    )
    db.add(material)
    await db.commit()
    await db.refresh(material)

    return SowUploadResponse(
        material_id=material.id,
        title=material.title,
        exam_type=material.exam_type,
        parsed_structure=parsed,
    )


@router.post("/sow/upload-image", response_model=SowUploadJobAccepted, status_code=202)
async def upload_sow_image(
    file: UploadFile = File(...),
    exam_type: str | None = Form(default=None),
    request: Request = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a SOW image. The file is read and validated synchronously
    so the 5 MB cap + content-type allowlist still hold before we
    return. OCR (Gemini Vision) + SOW AI parsing then run in a
    background task. The client polls
    `GET /api/v1/study/sow/jobs/{job_id}` to drive its progress bar
    through the 80→100 window.
    """
    content_length = request.headers.get("content-length") if request else None
    logger.info(
        "[sow/upload-image] HIT user=%s exam_type=%s filename=%s ctype=%s length=%s",
        current_user.id,
        exam_type,
        file.filename,
        file.content_type,
        content_length,
    )
    logger.info("[sow/upload-image] cloudinary_in_scope=%s", __name__)
    api_key = settings.gemini_api_key
    if not api_key:
        logger.error("[sow/upload-image] Gemini API key not configured")
        raise HTTPException(status_code=503, detail="Gemini not configured for image upload")

    # Content-type allowlist — `file.content_type` is set by the
    # client and is attacker-controlled. Reject anything other than
    # the types Gemini Vision actually accepts.
    ctype = (file.content_type or "").lower().split(";", 1)[0].strip()
    if ctype not in ALLOWED_IMAGE_TYPES:
        logger.warning("[sow/upload-image] unsupported ctype: %s", ctype)
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported image type: {ctype!r}. Use JPEG, PNG, WEBP, or HEIC.",
        )

    # Sanitize the filename BEFORE we read the body so a hostile
    # path-traversal filename never touches the DB or the log.
    safe_name = safe_filename(file.filename, fallback="upload.jpg")
    logger.info("[sow/upload-image] safe_name=%s", safe_name)

    try:
        contents = await file.read()
    except Exception as e:
        logger.exception("[sow/upload-image] file.read() failed: %s", e)
        raise HTTPException(status_code=400, detail=f"Could not read uploaded file: {e}")

    logger.info(
        "[sow/upload-image] read %s bytes (limit=%s)",
        len(contents),
        MAX_SOW_IMAGE_BYTES,
    )
    if not contents:
        logger.warning("[sow/upload-image] empty file")
        raise HTTPException(status_code=400, detail="Empty file")
    if len(contents) > MAX_SOW_IMAGE_BYTES:
        logger.warning(
            "[sow/upload-image] too large: %s > %s",
            len(contents),
            MAX_SOW_IMAGE_BYTES,
        )
        raise HTTPException(
            status_code=413,
            detail=f"Image too large (max {MAX_SOW_IMAGE_BYTES // (1024*1024)} MB)",
        )

    # Insert the job row, fire the worker, return immediately.
    job = SowUploadJob(id=str(uuid.uuid4()), user_id=current_user.id, status="queued")
    db.add(job)
    await db.commit()
    await db.refresh(job)

    logger.info(
        "[sow/upload-image] job queued id=%s user=%s dispatching background worker",
        job.id,
        current_user.id,
    )

    asyncio.create_task(
        _process_sow_image_job(
            job_id=job.id,
            user_id=current_user.id,
            exam_type=exam_type,
            contents=contents,
            safe_name=safe_name,
            content_type=file.content_type or "image/jpeg",
        )
    )

    return SowUploadJobAccepted(job_id=job.id, status="queued")


@router.post("/sow/upload-document", response_model=SowUploadJobAccepted, status_code=202)
async def upload_sow_document(
    file: UploadFile = File(...),
    exam_type: str | None = Form(default=None),
    request: Request = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a SOW document (PDF, DOCX, TXT). The file is read and
    validated synchronously so the 10 MB cap holds; text extraction
    + SOW AI parsing then run in a background task. The client
    polls `GET /api/v1/study/sow/jobs/{job_id}` to drive its progress
    bar through the 80→100 window.
    """
    content_length = request.headers.get("content-length") if request else None
    logger.info(
        "[sow/upload-document] HIT user=%s exam_type=%s filename=%s ctype=%s length=%s",
        current_user.id,
        exam_type,
        file.filename,
        file.content_type,
        content_length,
    )
    logger.info("[sow/upload-document] cloudinary_in_scope=%s", __name__)
    safe_name = safe_filename(file.filename, fallback="document")
    logger.info("[sow/upload-document] safe_name=%s", safe_name)

    try:
        contents = await file.read()
    except Exception as e:
        logger.exception(
            "[sow/upload-document] file.read() failed: %s", e,
        )
        raise HTTPException(status_code=400, detail=f"Could not read uploaded file: {e}")

    logger.info(
        "[sow/upload-document] read %s bytes (limit=%s)",
        len(contents),
        MAX_SOW_DOC_BYTES,
    )

    if not contents:
        logger.warning("[sow/upload-document] empty file")
        raise HTTPException(status_code=400, detail="Empty file")
    if len(contents) > MAX_SOW_DOC_BYTES:
        logger.warning(
            "[sow/upload-document] too large: %s > %s",
            len(contents),
            MAX_SOW_DOC_BYTES,
        )
        raise HTTPException(
            status_code=413,
            detail=f"Document too large (max {MAX_SOW_DOC_BYTES // (1024*1024)} MB)",
        )

    # Reject early if the extension is not one we know how to parse —
    # do this inline so the client gets a fast 400 instead of finding
    # out via the polling endpoint.
    filename_lower = safe_name.lower()
    if not (
        filename_lower.endswith(".pdf")
        or filename_lower.endswith(".docx")
        or filename_lower.endswith(".doc")
        or filename_lower.endswith(".txt")
    ):
        logger.warning(
            "[sow/upload-document] unsupported extension: %s", safe_name,
        )
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Upload PDF, DOCX, or TXT.",
        )

    job = SowUploadJob(id=str(uuid.uuid4()), user_id=current_user.id, status="queued")
    db.add(job)
    await db.commit()
    await db.refresh(job)

    logger.info(
        "[sow/upload-document] job queued id=%s user=%s dispatching background worker",
        job.id,
        current_user.id,
    )

    asyncio.create_task(
        _process_sow_document_job(
            job_id=job.id,
            user_id=current_user.id,
            exam_type=exam_type,
            contents=contents,
            safe_name=safe_name,
        )
    )

    return SowUploadJobAccepted(job_id=job.id, status="queued")


@router.get("/sow/jobs/{job_id}", response_model=SowUploadJobStatus)
async def get_sow_job_status(
    job_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Poll for the status of an in-flight SOW upload.

    Scoped to `current_user.id` so a user can never see another
    user's job state. Returns 404 on miss — we don't leak the
    existence of someone else's job.
    """
    row = await db.get(SowUploadJob, job_id)
    if row is None or row.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Job not found")
    return SowUploadJobStatus(
        job_id=row.id,
        status=row.status,
        material_id=row.material_id,
        error=row.error_message,
        updated_at=row.updated_at,
    )


# ── Background workers ───────────────────────────────────────────────


async def _ocr_image_with_gemini(
    contents: bytes, content_type: str, api_key: str
) -> str:
    """Send the image to Gemini Vision and return the extracted text.

    Kept as a private helper so the image worker stays small. Raises
    on any failure (network, non-200, empty body) so the worker can
    mark the job as `failed` with a category string instead of an
    exception that crashes the asyncio task.
    """
    import base64
    import httpx

    b64 = base64.b64encode(contents).decode()
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
    headers = {"x-goog-api-key": api_key}
    body = {
        "contents": [{
            "parts": [
                {"text": "Extract all text from this image. Return the raw text exactly as written, preserving structure (headings, bullet points, numbering)."},
                {"inline_data": {"mime_type": content_type, "data": b64}},
            ]
        }],
        "generationConfig": {"maxOutputTokens": 8000, "temperature": 0.1},
    }
    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(url, json=body, headers=headers)
        if resp.status_code != 200:
            raise RuntimeError(f"gemini_vision_http_{resp.status_code}")
        data = resp.json()
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError, TypeError) as exc:
        raise RuntimeError("gemini_vision_malformed_response") from exc


async def _extract_document_text(contents: bytes, filename: str) -> str:
    """Branch on extension and extract plain text from PDF / DOCX / TXT.

    Raises a `ValueError` with a user-safe category on any failure;
    the document worker translates that into the right
    `error_message` on the job row.
    """
    fname = filename.lower()
    if fname.endswith(".pdf"):
        try:
            import pypdf
            from io import BytesIO
            reader = pypdf.PdfReader(BytesIO(contents))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc)
            raise ValueError("pdf_extraction_failed") from exc
    if fname.endswith((".docx", ".doc")):
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(contents))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception as exc:
            logger.error("DOCX extraction failed: %s", exc)
            raise ValueError("docx_extraction_failed") from exc
    if fname.endswith(".txt"):
        try:
            return contents.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return contents.decode("latin-1")
            except Exception as exc:
                logger.error("Text file decode failed: %s", exc)
                raise ValueError("text_decode_failed") from exc
    raise ValueError("unsupported_file_type")


def _parse_ai_response(response_text: str) -> dict | None:
    """Best-effort JSON parse of the SOW parser's output.

    Some providers wrap JSON in a ```json ... ``` fenced code block;
    strip that before parsing. Returns None on failure.
    """
    import json as _json
    try:
        return _json.loads(response_text)
    except Exception:
        pass
    try:
        if "```json" in response_text:
            inner = response_text.split("```json", 1)[1]
            inner = inner.split("```", 1)[0]
            return _json.loads(inner)
    except Exception:
        pass
    try:
        if "```" in response_text:
            inner = response_text.split("```", 1)[1]
            inner = inner.split("```", 1)[0]
            return _json.loads(inner)
    except Exception:
        pass
    logger.error("SOW parser returned non-JSON: %s", response_text[:200])
    return None


async def _run_sow_parser(extracted_text: str) -> tuple[dict | None, str | None, str | None]:
    """Run the SOW parser AI prompt. Returns (parsed, provider, raw_response).

    Pulled out of the workers so the image + document paths share the
    exact same parsing logic.
    """
    prompt = SOW_PARSER.format(raw_text=extracted_text)
    ai_result = await route_ai(prompt, task_type="heavy")
    parsed = _parse_ai_response(ai_result.get("response", ""))
    return parsed, ai_result.get("provider"), ai_result.get("response", "")


async def _mark_job_status(
    job_id: str, status: str, *, material_id: int | None = None, error: str | None = None
) -> None:
    """Update a job row using its own session.

    The worker runs in a background task with no shared session, so
    each status flip opens `AsyncSessionLocal` briefly.
    """
    async with AsyncSessionLocal() as db:
        row = await db.get(SowUploadJob, job_id)
        if row is None:
            return
        row.status = status
        if material_id is not None:
            row.material_id = material_id
        if error is not None:
            row.error_message = error
        row.updated_at = datetime.utcnow()
        await db.commit()


async def _process_sow_image_job(
    job_id: str,
    user_id: int,
    exam_type: str | None,
    contents: bytes,
    safe_name: str,
    content_type: str,
) -> None:
    """Background worker for the image SOW upload.

    Opens its own `AsyncSessionLocal` so it doesn't share a session
    with the request that fired it. Never raises — every error path
    ends in a `failed` row so the polling endpoint has something to
    report.
    """
    logger.info(
        "[sow/upload-image worker] START job=%s user=%s bytes=%s cloudinary_available=%s",
        job_id, user_id, len(contents), "cloudinary" in dir(),
    )
    logger.info("[sow/upload-image worker] study_router_cloudinary_imports=%s", "cloudinary" in globals())
    try:
        await _mark_job_status(job_id, "processing")

        api_key = settings.gemini_api_key
        if not api_key:
            raise RuntimeError("gemini_not_configured")

        try:
            extracted_text = await _ocr_image_with_gemini(contents, content_type, api_key)
            logger.info(
                "[sow/upload-image worker] OCR done job=%s text_len=%s",
                job_id, len(extracted_text or ""),
            )
        except Exception as exc:
            logger.error("Gemini Vision OCR failed for job %s: %s", job_id, exc)
            raise RuntimeError("image_ocr_failed") from exc

        if not extracted_text.strip():
            raise RuntimeError("image_ocr_empty")

        try:
            parsed, provider, _ = await _run_sow_parser(extracted_text)
            logger.info(
                "[sow/upload-image worker] parser done job=%s provider=%s keys=%s",
                job_id, provider, list((parsed or {}).keys()) if isinstance(parsed, dict) else None,
            )
        except Exception as exc:
            logger.error("SOW parser AI failed for job %s: %s", job_id, exc)
            raise RuntimeError("ai_parser_failed") from exc

        title = (parsed or {}).get("title", safe_name) if isinstance(parsed, dict) else safe_name

        import json as _json
        async with AsyncSessionLocal() as db:
            material = StudyMaterial(
                user_id=user_id,
                title=title,
                exam_type=exam_type,
                raw_input=f"[IMAGE: {safe_name}]\n{extracted_text}",
                parsed_structure=_json.dumps(parsed) if parsed else None,
                original_file_data=contents,
                file_mime_type=file.content_type or "image/jpeg",
                ai_model_used=provider,
            )
            db.add(material)
            await db.commit()
            await db.refresh(material)
            material_id = material.id

        logger.info(
            "[sow/upload-image worker] DONE job=%s material_id=%s",
            job_id, material_id,
        )
        await _mark_job_status(job_id, "completed", material_id=material_id)
    except Exception as exc:
        logger.exception("SOW image job %s failed", job_id)
        await _mark_job_status(job_id, "failed", error=str(exc) or exc.__class__.__name__)


async def _process_sow_document_job(
    job_id: str,
    user_id: int,
    exam_type: str | None,
    contents: bytes,
    safe_name: str,
) -> None:
    """Background worker for the document SOW upload. Same shape as
    the image worker — runs extraction + AI parse, writes the
    material, updates the job row.
    """
    logger.info(
        "[sow/upload-document worker] START job=%s user=%s bytes=%s cloudinary_available=%s",
        job_id, user_id, len(contents), "cloudinary" in dir(),
    )
    logger.info("[sow/upload-document worker] study_router_cloudinary_imports=%s", "cloudinary" in globals())
    try:
        await _mark_job_status(job_id, "processing")

        try:
            extracted_text = await _extract_document_text(contents, safe_name)
            logger.info(
                "[sow/upload-document worker] extract done job=%s text_len=%s",
                job_id, len(extracted_text or ""),
            )
        except ValueError as exc:
            # Already a user-safe category from the extractor.
            logger.warning(
                "[sow/upload-document worker] extract ValueError job=%s err=%s",
                job_id, exc,
            )
            await _mark_job_status(job_id, "failed", error=str(exc))
            return
        except Exception as exc:
            logger.exception("Document extraction crashed for job %s: %s", job_id, exc)
            await _mark_job_status(job_id, "failed", error="document_extraction_failed")
            return

        if not extracted_text.strip():
            logger.warning(
                "[sow/upload-document worker] empty text job=%s",
                job_id,
            )
            await _mark_job_status(job_id, "failed", error="empty_extracted_text")
            return

        try:
            parsed, provider, _ = await _run_sow_parser(extracted_text)
            logger.info(
                "[sow/upload-document worker] parser done job=%s provider=%s",
                job_id, provider,
            )
        except Exception as exc:
            logger.exception("SOW parser AI failed for job %s: %s", job_id, exc)
            await _mark_job_status(job_id, "failed", error="ai_parser_failed")
            return

        title = (parsed or {}).get("title", safe_name) if isinstance(parsed, dict) else safe_name

        import json as _json
        async with AsyncSessionLocal() as db:
            filename_lower = safe_name.lower()
            if filename_lower.endswith(".pdf"):
                inferred_mime = "application/pdf"
            elif filename_lower.endswith((".docx", ".doc")):
                inferred_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif filename_lower.endswith(".txt"):
                inferred_mime = "text/plain"
            else:
                inferred_mime = "application/octet-stream"
            material = StudyMaterial(
                user_id=user_id,
                title=title,
                exam_type=exam_type,
                raw_input=f"[DOCUMENT: {safe_name}]\n{extracted_text}",
                parsed_structure=_json.dumps(parsed) if parsed else None,
                original_file_data=contents,
                file_mime_type=inferred_mime,
                ai_model_used=provider,
            )
            db.add(material)
            await db.commit()
            await db.refresh(material)
            material_id = material.id

        logger.info(
            "[sow/upload-document worker] DONE job=%s material_id=%s",
            job_id, material_id,
        )
        await _mark_job_status(job_id, "completed", material_id=material_id)
    except Exception as exc:
        logger.exception("SOW document job %s failed", job_id)
        await _mark_job_status(job_id, "failed", error=str(exc) or exc.__class__.__name__)


# ── GET /study/materials ────────────────────────────────────────────


@router.get("/materials", response_model=list[MaterialSummary])
async def list_materials(
    exam_type: str | None = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(StudyMaterial).where(StudyMaterial.user_id == current_user.id)
    if exam_type:
        query = query.where(or_(StudyMaterial.exam_type == exam_type, StudyMaterial.exam_type.is_(None)))
    result = await db.execute(query.order_by(StudyMaterial.created_at.desc()))
    materials = result.scalars().all()

    out = []
    for m in materials:
        asset_types = await _get_asset_types(db, m.id)
        out.append(MaterialSummary(
            id=m.id,
            title=m.title,
            exam_type=m.exam_type,
            asset_types=asset_types,
            created_at=m.created_at,
        ))
    return out


async def _get_asset_types(db: AsyncSession, material_id: int) -> list[str]:
    result = await db.execute(
        select(StudyAsset.asset_type)
        .where(StudyAsset.material_id == material_id)
        .distinct()
    )
    return [row[0] for row in result.all()]


# ── GET /study/materials/{id} ───────────────────────────────────────


@router.get("/materials/{material_id}", response_model=MaterialDetail)
async def get_material(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    assets_result = await db.execute(
        select(StudyAsset).where(StudyAsset.material_id == material_id)
    )
    assets = assets_result.scalars().all()

    # Get unlocked assets for this user
    unlocked_result = await db.execute(
        select(StudyTransaction.asset_id, StudyAsset.content_json)
        .join(StudyAsset, StudyTransaction.asset_id == StudyAsset.id)
        .where(
            StudyTransaction.user_id == current_user.id,
            StudyTransaction.reward_granted == True,
            StudyAsset.material_id == material_id,
        )
    )
    unlocked_data = {row[0]: row[1] for row in unlocked_result.all()}

    import json as _json
    parsed = _json.loads(material.parsed_structure) if material.parsed_structure else None

    asset_list = []
    for a in assets:
        asset_dict = {
            "id": a.id,
            "material_id": a.material_id,
            "type": a.asset_type,
            "points_to_unlock": a.points_to_unlock,
            "created_at": a.created_at.isoformat(),
        }
        
        # Include content if already unlocked
        if a.id in unlocked_data:
            asset_dict["unlocked"] = True
            asset_dict["content"] = _json.loads(unlocked_data[a.id])
        else:
            asset_dict["unlocked"] = False
            
        asset_list.append(asset_dict)

    return MaterialDetail(
        id=material.id,
        title=material.title,
        exam_type=material.exam_type,
        content=material.raw_input or None,
        parsed_structure=parsed,
        assets=asset_list,
        file_mime_type=material.file_mime_type or None,
        has_original_file=material.original_file_data is not None,
        created_at=material.created_at,
    )


# ── PATCH /study/materials/{id} ───────────────────────────────────────


@router.patch("/materials/{material_id}", response_model=MaterialDetail)
async def update_material(
    material_id: int,
    payload: MaterialUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    if payload.title is not None:
        material.title = payload.title.strip()
    if payload.exam_type is not None:
        material.exam_type = payload.exam_type.strip() or None

    await db.commit()
    await db.refresh(material)

    assets_result = await db.execute(
        select(StudyAsset).where(StudyAsset.material_id == material_id)
    )
    assets = assets_result.scalars().all()

    unlocked_result = await db.execute(
        select(StudyTransaction.asset_id, StudyAsset.content_json)
        .join(StudyAsset, StudyTransaction.asset_id == StudyAsset.id)
        .where(
            StudyTransaction.user_id == current_user.id,
            StudyTransaction.reward_granted == True,
            StudyAsset.material_id == material_id,
        )
    )
    unlocked_data = {row[0]: row[1] for row in unlocked_result.all()}

    import json as _json
    parsed = _json.loads(material.parsed_structure) if material.parsed_structure else None

    asset_list = []
    for a in assets:
        asset_dict = {
            "id": a.id,
            "material_id": a.material_id,
            "type": a.asset_type,
            "points_to_unlock": a.points_to_unlock,
            "created_at": a.created_at.isoformat(),
        }
        if a.id in unlocked_data:
            asset_dict["unlocked"] = True
            asset_dict["content"] = _json.loads(unlocked_data[a.id])
        else:
            asset_dict["unlocked"] = False
        asset_list.append(asset_dict)

    return MaterialDetail(
        id=material.id,
        title=material.title,
        exam_type=material.exam_type,
        content=material.raw_input or None,
        parsed_structure=parsed,
        assets=asset_list,
        file_mime_type=material.file_mime_type or None,
        has_original_file=material.original_file_data is not None,
        created_at=material.created_at,
    )


# ── DELETE /study/materials/{id} ──────────────────────────────────────


@router.delete("/materials/{material_id}", status_code=204)
async def delete_material(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    await db.delete(material)
    await db.commit()
    return Response(status_code=204)


# ── GET /study/materials/{id}/export ───────────────────────────────────


@router.get("/materials/{material_id}/export")
async def export_material(
    material_id: int,
    format: str = "pdf",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    logger.info("[export] material_id=%s format=%s user=%s", material_id, format, current_user.id)
    logger.info("[export] cloudinary_in_scope=%s", "cloudinary" in dir())
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    topic_names: list[str] = []
    if material.parsed_structure:
        import json as _json
        parsed = _json.loads(material.parsed_structure)
        topic_names = [t.get("name", "") for t in parsed.get("topics", []) if t.get("name")]

    assets_result = await db.execute(
        select(StudyAsset).where(StudyAsset.material_id == material_id)
    )
    assets = assets_result.scalars().all()

    content = material.raw_input or ""
    safe_title = material.title or "study-material"
    base_name = re.sub(r"^[A-Z]+ · ", "", safe_title).strip() or "study-material"
    base_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", base_name).strip() or "study_material"

    if format == "image":
        if material.original_file_data:
            return Response(
                content=material.original_file_data,
                media_type=material.file_mime_type or "application/octet-stream",
                headers={
                    "Content-Disposition": f'attachment; filename="{base_name}.png"',
                },
            )
        
        from PIL import Image, ImageDraw, ImageFont
        width, height = 800, 1000
        img = Image.new("RGB", (width, height), color="white")
        draw = ImageDraw.Draw(img)
        try:
            font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
            font_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
            font_meta = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
        except Exception:
            font_title = ImageFont.load_default()
            font_body = ImageFont.load_default()
            font_meta = ImageFont.load_default()

        y = 30
        draw.text((40, y), safe_title, fill="black", font=font_title)
        y += 50
        meta = f"Exam: {(material.exam_type or 'Custom').upper()}    Topics: {len(topic_names)}    Assets: {len(assets)}"
        draw.text((40, y), meta, fill="#666666", font=font_meta)
        y += 30
        draw.line([(40, y), (width - 40, y)], fill="#0E7C66", width=2)
        y += 20

        paragraphs = _split_paragraphs(content or "")
        for para in paragraphs:
            lines = _wrap_text(para, width - 80, font_body)
            for line in lines:
                if y > height - 40:
                    break
                draw.text((40, y), line, fill="black", font=font_body)
                y += 26
            y += 10

        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        return Response(
            content=buffer.getvalue(),
            media_type="image/png",
            headers={
                "Content-Disposition": f'attachment; filename="{base_name}.png"',
            },
        )

    if format == "pdf":
        from app.services.pdf_material import material_to_pdf
        pdf_bytes = material_to_pdf(
            title=safe_title,
            content=content,
            exam_type=material.exam_type,
            topic_names=topic_names,
            asset_count=len(assets),
            created_at=material.created_at,
            image_bytes=material.original_file_data,
        )
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{base_name}.pdf"',
            },
        )

    if format == "docx":
        import docx  # python-docx
        doc = docx.Document()
        doc.add_heading(safe_title, level=1)
        if material.exam_type:
            doc.add_paragraph(f"Exam: {material.exam_type.upper()}")
        doc.add_paragraph(f"Topics: {len(topic_names)}  |  Assets: {len(assets)}")
        doc.add_paragraph("")
        
        if material.original_file_data:
            try:
                img_stream = io.BytesIO(material.original_file_data)
                doc.add_picture(img_stream, width=docx.shared.Inches(5.5))
                doc.add_paragraph("")
            except Exception:
                pass
        
        if topic_names:
            doc.add_heading("Topics Covered", level=2)
            for idx, name in enumerate(topic_names, 1):
                doc.add_paragraph(f"{idx}. {name}", style="List Number")
            doc.add_paragraph("")
        doc.add_heading("Content", level=2)
        for para in (content or "").split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{base_name}.docx"',
            },
        )

    if format == "txt":
        body = "\n".join([
            safe_title,
            f"Exam: {(material.exam_type or 'Custom').upper()}",
            f"Topics: {len(topic_names)}",
            f"Assets: {len(assets)}",
            "",
            "---",
            "",
            content or "(No content available)",
        ])
        return Response(
            content=body.encode("utf-8"),
            media_type="text/plain",
            headers={
                "Content-Disposition": f'attachment; filename="{base_name}.txt"',
            },
        )

    raise HTTPException(status_code=400, detail="Unsupported export format. Use pdf, docx, txt, or image.")


# ── GET /study/materials/{id}/file ────────────────────────────────────


@router.get("/materials/{material_id}/file")
async def get_material_file(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    logger.info("[file] material_id=%s user=%s", material_id, current_user.id)
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")
    if not material.original_file_data:
        raise HTTPException(status_code=404, detail="Original file not available")

    mime = material.file_mime_type or "application/octet-stream"
    return Response(
        content=material.original_file_data,
        media_type=mime,
        headers={
            "Content-Disposition": f'inline; filename="material_{material_id}{_ext_from_mime(mime)}"',
            "Cache-Control": "private, max-age=3600",
        },
    )


# ── GET /study/materials/{id}/pages ───────────────────────────────────


@router.get("/materials/{material_id}/pages")
async def get_material_pages(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    logger.info("[pages] material_id=%s user=%s", material_id, current_user.id)
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")
    if not material.original_file_data:
        raise HTTPException(status_code=404, detail="Original file not available")

    mime = material.file_mime_type or ""
    if mime != "application/pdf":
        raise HTTPException(status_code=400, detail="Only PDF page rendering is supported")

    if len(material.original_file_data) == 0:
        raise HTTPException(status_code=400, detail="Empty PDF file")

    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.exception("[pages] PyMuPDF not installed")
        raise HTTPException(status_code=500, detail="PDF renderer not configured on server")

    try:
        pdf = fitz.open(stream=material.original_file_data, filetype="pdf")
    except Exception as exc:
        logger.exception("[pages] failed to open PDF material_id=%s err=%s", material_id, exc)
        raise HTTPException(status_code=400, detail="Stored file is not a valid PDF") from exc

    try:
        pages: list[dict] = []
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            pages.append({
                "page": page_num + 1,
                "total": len(pdf),
                "image_base64": __import__("base64").b64encode(img_bytes).decode("ascii"),
                "width": pix.width,
                "height": pix.height,
            })
        pdf.close()
        return {"pages": pages}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("[pages] PDF rendering failed material_id=%s err=%s", material_id, exc)
        raise HTTPException(status_code=500, detail=f"PDF rendering failed: {exc}") from exc


# ── POST /study/generate ────────────────────────────────────────────


@router.post("/generate", response_model=GenerateAssetResponse)
async def generate_asset(
    payload: GenerateAssetRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == payload.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    if not material.parsed_structure:
        raise HTTPException(status_code=400, detail="Material has no parsed structure. Re-upload or try again.")

    import json as _json
    parsed = _json.loads(material.parsed_structure)

    topics = parsed.get("topics", [])

    if payload.mode == "topic":
        if not payload.topic:
            raise HTTPException(status_code=400, detail="Topic name is required in topic mode.")
        topic_data = None
        for t in topics:
            if t.get("name") == payload.topic:
                topic_data = t
                break
        if topic_data is None:
            available = [t.get("name", "") for t in topics]
            raise HTTPException(
                status_code=404,
                detail=f"Topic '{payload.topic}' not found. Available topics: {', '.join(available)}",
            )
        context_parts = []
        context_parts.append(f"Topic: {topic_data['name']}")
        for st in topic_data.get("subtopics", []):
            context_parts.append(f"  - {st}")
        for cc in topic_data.get("key_concepts", []):
            context_parts.append(f"    * {cc}")
        context = "\n".join(context_parts)

        if payload.asset_type == "mcq":
            prompt = MCQ_TOPIC_GENERATOR.format(topic=payload.topic, context=context, count=payload.count, difficulty=payload.difficulty)
        elif payload.asset_type == "flashcard":
            prompt = FLASHCARD_TOPIC_GENERATOR.format(topic=payload.topic, context=context, count=payload.count)
        elif payload.asset_type == "essay":
            prompt = ESSAY_TOPIC_GENERATOR.format(topic=payload.topic, context=context, count=payload.count)
        elif payload.asset_type == "diagram":
            prompt = DIAGRAM_GENERATOR.format(topic=payload.topic, context=context, education_level=payload.education_level or "secondary")
        elif payload.asset_type == "video":
            prompt = VIDEO_SCRIPT_GENERATOR.format(topic=payload.topic, context=context, education_level=payload.education_level or "secondary")
        else:
            raise HTTPException(status_code=400, detail=f"Unknown asset type: {payload.asset_type}")

    else:
        context_parts = []
        for topic in topics:
            context_parts.append(f"Topic: {topic['name']}")
            for st in topic.get("subtopics", []):
                context_parts.append(f"  - {st}")
            for cc in topic.get("key_concepts", []):
                context_parts.append(f"    * {cc}")
        context = "\n".join(context_parts)

        if payload.asset_type == "mcq":
            prompt = MCQ_ALL_TOPICS_GENERATOR.format(context=context, count=payload.count, difficulty=payload.difficulty)
        elif payload.asset_type == "flashcard":
            prompt = FLASHCARD_ALL_TOPICS_GENERATOR.format(context=context, count=payload.count)
        elif payload.asset_type == "essay":
            prompt = ESSAY_ALL_TOPICS_GENERATOR.format(context=context, count=payload.count)
        elif payload.asset_type == "diagram":
            prompt = DIAGRAM_GENERATOR.format(topic="All Topics", context=context, education_level=payload.education_level or "secondary")
        elif payload.asset_type == "video":
            prompt = VIDEO_SCRIPT_GENERATOR.format(topic="All Topics", context=context, education_level=payload.education_level or "secondary")
        else:
            raise HTTPException(status_code=400, detail=f"Unknown asset type: {payload.asset_type}")

    ai_result = await route_ai(prompt, task_type="fast", db=db)

    content = None
    try:
        content = _json.loads(ai_result["response"])
    except Exception:
        logger.error("Asset generator returned non-JSON: %s", ai_result["response"][:200])
        raise HTTPException(status_code=502, detail="AI returned invalid format. Try again.")

    points_cost = VIDEO_UNLOCK_POINTS_COST if payload.asset_type == "video" else UNLOCK_POINTS_COST

    asset = StudyAsset(
        material_id=payload.material_id,
        asset_type=payload.asset_type,
        content_json=_json.dumps(content),
        points_to_unlock=points_cost,
    )
    db.add(asset)
    await db.commit()
    await db.refresh(asset)

    return GenerateAssetResponse(assets=[content])


# ── POST /study/progress ─────────────────────────────────────────────
@router.post("/progress", response_model=StudyProgressResponse, status_code=201)
async def update_progress(
    payload: StudyProgressUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == payload.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    progress = StudyProgress(
        user_id=current_user.id,
        material_id=payload.material_id,
        topic_index=payload.topic_index,
        topic_name=payload.topic_name,
        status=payload.status,
        mastery_score=payload.mastery_score,
    )
    db.add(progress)
    await db.commit()
    await db.refresh(progress)
    return progress


@router.get("/materials/{material_id}/progress", response_model=StudyProgressListResponse)
async def get_progress(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    import json as _json
    parsed = _json.loads(material.parsed_structure or '{"topics": []}')
    total_topics = len(parsed.get("topics", []))

    progress_result = await db.execute(
        select(StudyProgress).where(
            StudyProgress.user_id == current_user.id,
            StudyProgress.material_id == material_id,
        )
    )
    progress_rows = progress_result.scalars().all()

    mastered = sum(1 for p in progress_rows if p.status == "mastered")
    reviewing = sum(1 for p in progress_rows if p.status == "reviewing")
    not_started = total_topics - mastered - reviewing

    return StudyProgressListResponse(
        material_id=material_id,
        total_topics=total_topics,
        mastered=mastered,
        reviewing=reviewing,
        not_started=max(0, not_started),
        progress=[
            StudyProgressResponse(
                id=p.id,
                material_id=p.material_id,
                topic_index=p.topic_index,
                topic_name=p.topic_name,
                status=p.status,
                mastery_score=p.mastery_score,
                last_reviewed_at=p.last_reviewed_at,
            )
            for p in progress_rows
        ],
    )


# ── POST /study/examples/generate ───────────────────────────────────


@router.post("/examples/generate", response_model=GenerateAssetResponse, status_code=201)
async def generate_example(
    payload: ExampleGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == payload.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    if not material.parsed_structure:
        raise HTTPException(status_code=400, detail="Material has no parsed structure. Re-upload or try again.")

    import json as _json
    parsed = _json.loads(material.parsed_structure)
    topics = parsed.get("topics", [])

    context_parts = []
    target_topic = payload.topic
    if payload.mode == "topic" and target_topic:
        topic_data = next((t for t in topics if t.get("name") == target_topic), None)
        if not topic_data:
            available = [t.get("name", "") for t in topics]
            raise HTTPException(status_code=404, detail=f"Topic '{target_topic}' not found. Available: {', '.join(available)}")
        context_parts.append(f"Topic: {topic_data['name']}")
        for st in topic_data.get("subtopics", []):
            context_parts.append(f"  - {st}")
        for cc in topic_data.get("key_concepts", []):
            context_parts.append(f"    * {cc}")
    else:
        for topic in topics:
            context_parts.append(f"Topic: {topic['name']}")
            for st in topic.get("subtopics", []):
                context_parts.append(f"  - {st}")
            for cc in topic.get("key_concepts", []):
                context_parts.append(f"    * {cc}")
        target_topic = "All Topics"

    context = "\n".join(context_parts)
    prompt = EXAMPLE_GENERATOR.format(
        topic=target_topic,
        context=context,
        education_level=payload.education_level or "secondary",
        subject_hints=payload.subject_hints,
    )

    ai_result = await route_ai(prompt, task_type="fast", db=db)

    content = None
    try:
        content = _json.loads(ai_result["response"])
    except Exception:
        logger.error("Example generator returned non-JSON: %s", ai_result["response"][:200])
        raise HTTPException(status_code=502, detail="AI returned invalid format. Try again.")

    asset = StudyAsset(
        material_id=payload.material_id,
        asset_type="example",
        content_json=_json.dumps(content),
        points_to_unlock=UNLOCK_POINTS_COST,
    )
    db.add(asset)
    await db.commit()
    await db.refresh(asset)

    return GenerateAssetResponse(assets=[content])


# ── POST /study/examples/check ──────────────────────────────────────


@router.post("/examples/check", response_model=ExampleCheckResponse)
async def check_example(
    payload: ExampleCheckRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    asset_result = await db.execute(
        select(StudyAsset).where(
            StudyAsset.id == payload.example_id,
            StudyAsset.asset_type == "example",
        )
    )
    asset = asset_result.scalar_one_or_none()
    if not asset:
        raise HTTPException(status_code=404, detail="Example not found")

    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == asset.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    if not material_result.scalar_one_or_none():
        raise HTTPException(status_code=403, detail="Not your material")

    import json as _json
    content = _json.loads(asset.content_json)
    try_data = content.get("try_yourself", {})

    checker_prompt = f"""You are a friendly tutor checking a student's answer to a practice problem.

Problem: {try_data.get('problem', '')}
Correct final answer: {try_data.get('final_answer', '')}
Student's answer: {payload.user_answer}

Available hints (use only if the student is wrong):
{chr(10).join(f"- {h}" for h in try_data.get('hints', []))}

Solution steps:
{chr(10).join(f"{i+1}. {s}" for i, s in enumerate(try_data.get('solution_steps', [])))}

Output strict JSON only. No markdown. No backticks. No extra text.
{{
  "correct": true/false,
  "feedback": "Encouraging feedback on their answer",
  "hint": "A helpful hint if they got it wrong (null if correct)",
  "next_step_instruction": "What to do next",
  "show_answer": false
}}

Rules:
- Be encouraging, not discouraging
- If correct: congratulate and suggest moving to the next topic
- If wrong: explain what went wrong in simple terms, don't just say "incorrect"
- Use the hints provided above
- show_answer should be true only if the student has attempted 3+ times (we'll track this client-side for now)"""

    ai_result = await route_ai(checker_prompt, task_type="fast", db=db)

    check_result = None
    try:
        check_result = _json.loads(ai_result["response"])
    except Exception:
        logger.error("Example checker returned non-JSON: %s", ai_result["response"][:200])
        raise HTTPException(status_code=502, detail="AI returned invalid format. Try again.")

    return ExampleCheckResponse(
        correct=check_result.get("correct", False),
        feedback=check_result.get("feedback", "Keep trying!"),
        hint=check_result.get("hint"),
        next_step_instruction=check_result.get("next_step_instruction"),
        show_answer=check_result.get("show_answer", False),
    )


# ── POST /study/chat (streaming) ────────────────────────────────────


@router.post("/chat")
async def chat_study(
    payload: ChatRequest,
    education_level: str | None = None,
    difficulty: str = "medium",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == payload.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    context = ""
    if material.parsed_structure:
        import json as _json
        parsed = _json.loads(material.parsed_structure)
        context_parts = []
        for topic in parsed.get("topics", []):
            context_parts.append(f"Topic: {topic['name']}")
            for st in topic.get("subtopics", []):
                context_parts.append(f"  - {st}")
        context = "\n".join(context_parts)

    system_prompt = CHAT_TUTOR_SYSTEM.format(
        context=context or "No structured context available.",
        education_level=education_level or "secondary",
        difficulty=difficulty,
    )
    full_prompt = f"{system_prompt}\n\nStudent question: {payload.message}"

    async def generate():
        ai_result = await route_ai(full_prompt, task_type="chat", max_tokens=2000, db=db)
        text = ai_result["response"]
        # Stream token-by-token in small chunks for the frontend
        words = text.split(" ")
        chunk = []
        for word in words:
            chunk.append(word)
            yield " ".join(chunk) + " "
            chunk = []
        if chunk:
            yield " ".join(chunk)

    return StreamingResponse(generate(), media_type="text/plain")


# ── POST /study/unlock ──────────────────────────────────────────────


@router.post("/unlock", response_model=UnlockResponse)
async def unlock_asset(
    payload: UnlockRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Check if user is premium (imports at top needed)
    from app.services.subscription import is_premium
    
    asset_result = await db.execute(
        select(StudyAsset).where(StudyAsset.id == payload.asset_id)
    )
    asset = asset_result.scalar_one_or_none()
    if not asset:
        raise HTTPException(status_code=404, detail="Asset not found")

    # Verify ownership: user must own the parent material
    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == asset.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    if not material_result.scalar_one_or_none():
        raise HTTPException(status_code=403, detail="Not your material")

    # Premium users unlock for free
    if is_premium(current_user):
        import json as _json
        txn = StudyTransaction(
            user_id=current_user.id,
            asset_id=asset.id,
            method="premium",
            points_spent=0,
            reward_granted=True,
        )
        db.add(txn)
        await db.commit()
        return UnlockResponse(
            unlocked=True,
            content=_json.loads(asset.content_json),
            new_balance=current_user.points_balance,
            method="premium",
            points_spent=0,
        )

    if payload.method == "points":
        if settings.wallet_split_enabled:
            balance_result = await db.execute(
                select(User.service_credit_balance).where(User.id == current_user.id)
            )
            balance = balance_result.scalar_one() or 0
            if balance < asset.points_to_unlock:
                raise HTTPException(
                    status_code=402,
                    detail=f"Need {asset.points_to_unlock} service credits. You have {balance}.",
                )
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(service_credit_balance=User.service_credit_balance - asset.points_to_unlock)
            )
        else:
            balance_result = await db.execute(
                select(User.points_balance).where(User.id == current_user.id)
            )
            balance = balance_result.scalar_one() or 0
            if balance < asset.points_to_unlock:
                raise HTTPException(
                    status_code=402,
                    detail=f"Need {asset.points_to_unlock} pts. You have {balance}.",
                )
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(points_balance=User.points_balance - asset.points_to_unlock)
            )

        txn = StudyTransaction(
            user_id=current_user.id,
            asset_id=asset.id,
            method="points",
            points_spent=asset.points_to_unlock,
            reward_granted=True,
        )
        db.add(txn)
        await db.commit()

        if settings.wallet_split_enabled:
            new_balance_result = await db.execute(
                select(User.service_credit_balance).where(User.id == current_user.id)
            )
        else:
            new_balance_result = await db.execute(
                select(User.points_balance).where(User.id == current_user.id)
            )
        new_balance = new_balance_result.scalar_one() or 0

        import json as _json
        return UnlockResponse(
            unlocked=True,
            content=_json.loads(asset.content_json),
            new_balance=new_balance,
            method="points",
            points_spent=asset.points_to_unlock,
        )

    raise HTTPException(status_code=400, detail="Invalid method")


# ── POST /study/materials/{material_id}/unlock-audio ───────────────────

AUDIO_UNLOCK_BASE_SV = 20
AUDIO_UNLOCK_CHAR_STEP = 500


def _audio_unlock_cost_sv(content: str | None) -> int:
    length = len(content or "")
    return AUDIO_UNLOCK_BASE_SV + max(0, (length // AUDIO_UNLOCK_CHAR_STEP) * 10)


class AudioUnlockResponse(BaseModel):
    unlocked: bool
    material_id: int
    cost_sv: int
    method: str
    new_balance: int
    url: str | None = None
    provider: str | None = None


@router.post("/materials/{material_id}/unlock-audio", response_model=AudioUnlockResponse)
async def unlock_material_audio(
    material_id: int,
    payload: dict,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    method = str(payload.get("method") or "sv").lower()
    if method not in {"sv", "ad"}:
        raise HTTPException(status_code=400, detail="method must be 'sv' or 'ad'")

    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    material = material_result.scalar_one_or_none()
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")

    existing_result = await db.execute(
        select(AudioUnlock).where(
            AudioUnlock.user_id == current_user.id,
            AudioUnlock.material_id == material_id,
        )
    )
    if existing_result.scalar_one_or_none():
        return AudioUnlockResponse(
            unlocked=True,
            material_id=material_id,
            cost_sv=0,
            method="existing",
            new_balance=(
                current_user.service_credit_balance
                if settings.wallet_split_enabled
                else current_user.points_balance
            ),
            url=f"/api/v1/study/tts/audio/{material_id}_cached.mp3",
            provider="cached",
        )

    cost_sv = _audio_unlock_cost_sv(material.content)

    if method == "sv":
        if settings.wallet_split_enabled:
            balance_result = await db.execute(
                select(User.service_credit_balance).where(User.id == current_user.id)
            )
            balance = balance_result.scalar_one() or 0
            if balance < cost_sv:
                raise HTTPException(
                    status_code=402,
                    detail=f"Need {cost_sv} sv to unlock audio. You have {balance}.",
                )
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(service_credit_balance=User.service_credit_balance - cost_sv)
            )
        else:
            balance_result = await db.execute(
                select(User.points_balance).where(User.id == current_user.id)
            )
            balance = balance_result.scalar_one() or 0
            if balance < cost_sv:
                raise HTTPException(
                    status_code=402,
                    detail=f"Need {cost_sv} pts to unlock audio. You have {balance}.",
                )
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(points_balance=User.points_balance - cost_sv)
            )

        unlock = AudioUnlock(
            user_id=current_user.id,
            material_id=material_id,
            method="sv",
            cost_sv=cost_sv,
        )
        db.add(unlock)
        await db.commit()

        new_balance = (
            current_user.service_credit_balance
            if settings.wallet_split_enabled
            else current_user.points_balance
        )
        return AudioUnlockResponse(
            unlocked=True,
            material_id=material_id,
            cost_sv=cost_sv,
            method="sv",
            new_balance=new_balance,
            url=f"/api/v1/study/tts/audio/{material_id}_cached.mp3",
            provider="unlocked",
        )



class AudioUnlockStatusResponse(BaseModel):
    unlocked: bool
    material_id: int
    method: str | None = None
    cost_sv: int = 0


@router.get("/materials/{material_id}/audio-status", response_model=AudioUnlockStatusResponse)
async def get_audio_unlock_status(
    material_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    if not material_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Material not found")

    from app.services.subscription import is_premium
    if is_premium(current_user):
        return AudioUnlockStatusResponse(unlocked=True, material_id=material_id, method="premium")

    unlock_result = await db.execute(
        select(AudioUnlock).where(
            AudioUnlock.user_id == current_user.id,
            AudioUnlock.material_id == material_id,
        )
    )
    unlock = unlock_result.scalar_one_or_none()
    if unlock:
        return AudioUnlockStatusResponse(unlocked=True, material_id=material_id, method=unlock.method, cost_sv=unlock.cost_sv)

    return AudioUnlockStatusResponse(unlocked=False, material_id=material_id, cost_sv=_audio_unlock_cost_sv(material.content))


# ── POST /study/quiz/complete ────────────────────────────────────────


BONUS_THRESHOLD = 80
BONUS_POINTS = 20


@router.post("/quiz/complete", response_model=QuizCompleteResponse)
async def complete_quiz(
    payload: QuizCompleteRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    asset_result = await db.execute(
        select(StudyAsset).where(StudyAsset.id == payload.asset_id)
    )
    asset = asset_result.scalar_one_or_none()
    if not asset:
        raise HTTPException(status_code=404, detail="Asset not found")

    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == asset.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    if not material_result.scalar_one_or_none():
        raise HTTPException(status_code=403, detail="Not your material")

    bonus_awarded = False
    bonus_points = 0
    new_balance = current_user.service_credit_balance if settings.wallet_split_enabled else current_user.points_balance

    if payload.score >= BONUS_THRESHOLD:
        bonus_points = BONUS_POINTS
        if settings.wallet_split_enabled:
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(service_credit_balance=User.service_credit_balance + bonus_points)
            )
        else:
            await db.execute(
                update(User)
                .where(User.id == current_user.id)
                .values(points_balance=User.points_balance + bonus_points)
            )
        txn = StudyTransaction(
            user_id=current_user.id,
            asset_id=asset.id,
            method="quiz_bonus",
            points_spent=-bonus_points,
            reward_granted=True,
        )
        db.add(txn)
        await db.commit()
        if settings.wallet_split_enabled:
            new_balance_result = await db.execute(
                select(User.service_credit_balance).where(User.id == current_user.id)
            )
        else:
            new_balance_result = await db.execute(
                select(User.points_balance).where(User.id == current_user.id)
            )
        new_balance = new_balance_result.scalar_one() or 0
        bonus_awarded = True

    # Send push notification for bonus (fire-and-forget)
    if bonus_awarded:
        try:
            from app.services.fcm import send_push_notification_background
            from app.services.notifications import create_notification_background
            import asyncio
            asyncio.create_task(send_push_notification_background(
                user_id=current_user.id,
                title="Quiz Bonus! 🎉",
                body=f"You scored {payload.score}% and earned +{bonus_points} points!",
                data={"type": "quiz_bonus", "score": str(payload.score), "points": str(bonus_points)},
                category="study_reminders",
            ))
            asyncio.create_task(create_notification_background(
                user_id=current_user.id,
                title="Quiz Bonus! 🎉",
                body=f"You scored {payload.score}% and earned +{bonus_points} points!",
                category="study_reminders",
                data={"type": "quiz_bonus", "score": payload.score, "points": bonus_points},
            ))
        except Exception:
            pass

    return QuizCompleteResponse(
        bonus_awarded=bonus_awarded,
        bonus_points=bonus_points,
        new_balance=new_balance,
        message=(
            f"Great job! +{bonus_points} pts for scoring {payload.score}%"
            if bonus_awarded
            else f"Score: {payload.score}%. Get {BONUS_THRESHOLD}%+ for a +{BONUS_POINTS} pts bonus!"
        ),
    )


# ── POST /study/session/start ───────────────────────────────────────


from pydantic import BaseModel


class SessionStartRequest(BaseModel):
    material_id: int


class SessionStartResponse(BaseModel):
    session_id: int
    started_at: str


class SessionEndRequest(BaseModel):
    session_id: int


class SessionEndResponse(BaseModel):
    session_id: int
    duration_seconds: int
    ended_at: str


class StudySession:
    """DB-backed session tracking using ReadingSession."""

    @classmethod
    async def start(cls, db: AsyncSession, user_id: int, material_id: int) -> int:
        session = ReadingSession(
            user_id=user_id,
            content_id=material_id,
            start_time=datetime.utcnow(),
            duration_seconds=0,
            points_earned=0,
            verified=False,
            scroll_events=0,
            total_paused_seconds=0,
        )
        db.add(session)
        await db.commit()
        await db.refresh(session)
        return session.id

    @classmethod
    async def end(cls, db: AsyncSession, session_id: int, user_id: int) -> dict | None:
        result = await db.execute(
            select(ReadingSession).where(
                ReadingSession.id == session_id,
                ReadingSession.user_id == user_id,
                ReadingSession.end_time.is_(None),
            )
        )
        session = result.scalar_one_or_none()
        if not session:
            return None

        ended_at = datetime.utcnow()
        duration = int((ended_at - session.start_time).total_seconds())
        session.end_time = ended_at
        session.duration_seconds = duration
        await db.commit()
        await db.refresh(session)

        return {
            "session_id": session.id,
            "user_id": session.user_id,
            "material_id": session.content_id,
            "started_at": session.start_time,
            "ended_at": session.end_time,
            "duration_seconds": session.duration_seconds,
        }

    @classmethod
    async def get(cls, db: AsyncSession, session_id: int, user_id: int) -> dict | None:
        result = await db.execute(
            select(ReadingSession).where(
                ReadingSession.id == session_id,
                ReadingSession.user_id == user_id,
            )
        )
        session = result.scalar_one_or_none()
        if not session:
            return None
        return {
            "session_id": session.id,
            "user_id": session.user_id,
            "material_id": session.content_id,
            "started_at": session.start_time,
            "ended_at": session.end_time,
            "duration_seconds": session.duration_seconds,
        }


@router.post("/session/start", response_model=SessionStartResponse, status_code=201)
async def start_study_session(
    payload: SessionStartRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Start a study session for time tracking"""
    # Verify material ownership
    material_result = await db.execute(
        select(StudyMaterial).where(
            StudyMaterial.id == payload.material_id,
            StudyMaterial.user_id == current_user.id,
        )
    )
    if not material_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Material not found")
    
    session_id = await StudySession.start(db, current_user.id, payload.material_id)
    
    return SessionStartResponse(
        session_id=session_id,
        started_at=datetime.utcnow().isoformat(),
    )


@router.post("/session/end", response_model=SessionEndResponse)
async def end_study_session(
    payload: SessionEndRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """End a study session and get duration"""
    session = await StudySession.end(db, payload.session_id, current_user.id)
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if session["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your session")
    
    # Notify on study milestone: 30+ minutes
    if session.get("duration_seconds", 0) >= 30 * 60:
        try:
            from app.services.fcm import send_push_notification_background
            from app.services.notifications import create_notification_background
            import asyncio
            duration_min = session["duration_seconds"] // 60
            asyncio.create_task(send_push_notification_background(
                user_id=current_user.id,
                title="Study Milestone! 📚",
                body=f"You studied for {duration_min} minutes. Keep it up!",
                data={"type": "study_reminder", "duration_minutes": str(duration_min)},
                category="study_reminders",
            ))
            asyncio.create_task(create_notification_background(
                user_id=current_user.id,
                title="Study Milestone! 📚",
                body=f"You studied for {duration_min} minutes. Keep it up!",
                category="study_reminders",
                data={"type": "study_reminder", "duration_minutes": duration_min},
            ))
        except Exception:
            pass

    return SessionEndResponse(
        session_id=payload.session_id,
        duration_seconds=session["duration_seconds"],
        ended_at=session["ended_at"].isoformat(),
    )


# ── POST /study/tts ───────────────────────────────────────────────────
class TtsResponse(BaseModel):
    url: str
    provider: str
    cached: bool


@router.post("/tts", response_model=TtsResponse)
async def synthesize_study_tts(
    payload: dict,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """On-demand TTS for study material text.

    Body: { "text": str, "voice"?: str, "provider"?: str, "material_id"?: int }
    Returns: { "url": str, "provider": str, "cached": bool }

    Uses the hybrid provider chain in services/tts_hybrid.py.
    Audio is cached to disk; repeat reads return the cached URL.
    Gated on premium or audio_unlock for the material.
    """
    text = payload.get("text") or ""
    if not text.strip():
        raise HTTPException(status_code=400, detail="text is required")
    if len(text) > 4000:
        text = text[:4000]

    material_id = payload.get("material_id")
    if material_id is not None:
        from app.services.subscription import is_premium
        if not is_premium(current_user):
            unlock_result = await db.execute(
                select(AudioUnlock).where(
                    AudioUnlock.user_id == current_user.id,
                    AudioUnlock.material_id == int(material_id),
                )
            )
            if not unlock_result.scalar_one_or_none():
                raise HTTPException(
                    status_code=403,
                    detail="Audio locked. Unlock this material to listen.",
                )

    voice = payload.get("voice") or settings.tts_default_voice
    provider = payload.get("provider") or settings.tts_default_provider

    try:
        audio, used_provider = await tts_hybrid.synthesize_study_audio(text, voice, provider=provider)
    except Exception as exc:
        logger.error("TTS synthesis failed for user=%s: %s", current_user.id, exc)
        raise HTTPException(status_code=502, detail=f"TTS synthesis failed: {exc}")

    cache = tts_hybrid._cache_path(text, voice, used_provider)
    return TtsResponse(
        url=f"/api/v1/study/tts/audio/{cache.name}",
        provider=used_provider,
        cached=True,
    )


# ── GET /study/tts/audio/{filename} ──────────────────────────────────
@router.get("/tts/audio/{filename}")
async def serve_study_tts_audio(filename: str):
    """Serve cached TTS audio files. Public, cached 1 day."""
    safe_name = filename.replace("..", "_").replace("/", "_")
    candidate = AUDIO_CACHE_DIR / "tts" / "*" / safe_name
    matches = list(AUDIO_CACHE_DIR.glob(f"tts/*/{safe_name}"))
    if not matches:
        raise HTTPException(status_code=404, detail="Audio not found")
    path = matches[0]
    return FileResponse(
        path,
        media_type="audio/mpeg",
        headers={"Cache-Control": "public, max-age=86400"},
    )


# ── Export helpers ─────────────────────────────────────────────────────

def _split_paragraphs(text: str) -> list[str]:
    parts = text.split("\n\n")
    return [p.strip() for p in parts if p.strip()]


def _wrap_text(text: str, max_width: int, font) -> list[str]:
    """Very basic word-wrap for PIL bitmap fonts."""
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        try:
            bbox = font.getbbox(test)
            width = bbox[2] - bbox[0]
        except Exception:
            width = len(test) * 10
        if width > max_width and current:
            lines.append(current)
            current = word
        else:
            current = test
    if current:
        lines.append(current)
    return lines or [""]


def _ext_from_mime(mime: str) -> str:
    mapping = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/webp": ".webp",
        "image/heic": ".heic",
        "application/pdf": ".pdf",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt",
    }
    return mapping.get(mime.lower().split(";", 1)[0].strip(), "")

